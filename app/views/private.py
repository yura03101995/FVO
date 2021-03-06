#-.- coding: utf-8 -.-
from app import app, db
from flask_login import current_user
from sqlalchemy import text
from sqlalchemy.sql import select, and_
from docx import Document as Doc

from openpyxl import load_workbook

from app.models import User, VUS, Document, Student_info, Basic_information, Comments
from app.models import Certificates_change_name, Communications, Passports, International_passports
from app.models import Registration_certificates, Middle_education, Spec_middle_education
from app.models import High_education, Military_education, Languages, Mothers_fathers
from app.models import Brothers_sisters_children, Married_certificates, Personal_data, Admins_vuses
from app.models.easy import *

from werkzeug.security import generate_password_hash
from flask import request, send_from_directory
import datetime
import json
from app.views.easy import *
from app.models.easy import *
from app.config import USER_PATH
import os
import random, string
import re
from keywords import *
from transliteration import *
from zipfile import ZipFile, ZIP_DEFLATED
import sys

def create_account(login, password, userData):
    hash = generate_password_hash(password)

    new_user = User(login = login, password = hash, entrance_year = int(userData['year']))
    new_user.vus_id = userData['vus'].id
    
    student_info = Student_info()

    new_user.students_info = student_info

    # add basic information
    basic_information = Basic_information(last_name=userData['lastName'], first_name=userData['firstName'], 
                                          middle_name=userData['middleName'])
    student_info['basic_information'] = basic_information
    db.session.add(basic_information)

    # add comments 
    comments = Comments()
    student_info['comments'] = comments
    db.session.add(comments)

    # common interface tables
    for table in (get_user_tables() + get_admin_tables()):
        if table != 'basic_information':
            section = get_class_by_tablename(table)()
            if section.is_fixed:
                student_info[table] = section
                student_info['table_'+table] = TABLE_STATES['NOT_EDITED']
                db.session.add(section)
    
    db.session.add(new_user)
    db.session.add(student_info)
    db.session.commit()
    return True

def create_admin_account(data):
    hash = generate_password_hash(data['password'])

    new_user = User(login = data['login'], 
                    password = hash,
                    vus_id = int( data['vus_id'] ),
                    role = USER_STATES[ data['role'] ]
                    )

    db.session.add(new_user)
    db.session.commit()
    return True

@app.route('/comment_user', methods=['POST'])
def comment_user():
    data = request.form
    user_id = data['id'];
    comment = data['comment']
    comment_user = Comments.query.get(user_id);
    if(comment_user):
        comment_user.comment = unicode(comment);
    else:
        comment_user = Comments(
            id = user_id,
            comment = unicode(comment)
        )
    db.session.add(comment_user)
    db.session.commit()
    return gen_success()

# Rule admins
@app.route('/make_account', methods=['POST'])
def make_account():
    data = request.form
    if 'login' not in data or 'password' not in data:
        return gen_error('Wrong data sent to server (must be [login, password]).')
    user = User.query.filter_by(login=data['login']).first()
    if user:
        return gen_success(message={'status':'error', 'error' : u'Пользователь с таким логином уже существует'})
    create_admin_account(data)
    return gen_success(message={'status':'ok'})

@app.route('/delete_admin_account', methods=['POST'])
def delete_admin_account():
    data = request.form
    user = User.query.filter_by(id=data['id'])
    if user:
        user.delete()
        db.session.commit()
        return gen_success(message={'status':'ok'})
    else:
        return gen_success(message={'status':'error', 'error':u'Данный администратор уже удален'})

@app.route('/change_admin_pswd', methods=['POST'])
def change_admin_pswd():
    data = request.form
    user = User.query.get(data['id'])
    if user:
        if data['new_psw'] != u'':
            user.password = generate_password_hash(data['new_psw'])
            db.session.commit()
            return gen_success(message={'status':'ok'})
        else:
            return gen_success(message={'status':'error','error':u'Пароль не может быть пустым'})
    else:
        return gen_success(message={'status':'error','error':u'Ошибка обращения к администратору'})

@app.route('/add_new_admin', methods=['POST'])
def add_new_admin():
    print 'begin'
    data = request.form
    vus_for_write = set(json.loads(data['vus_for_write']))
    vus_for_read = set(json.loads(data['vus_for_read']))
    print vus_for_read
    if vus_for_write & vus_for_read:
        return gen_success(message={'status':'error', 'error':u'Есть пересечения между ВУС на запись и ВУС на чтение'})
    user = User();
    user.login = data['login']
    user.password = generate_password_hash(data['password'])
    user.role = USER_STATES['ROLE_ADMIN']
    db.session.add(user)
    db.session.flush()
    for vus in vus_for_write:
        admin_vus = Admins_vuses()
        admin_vus.user_id = user.id
        admin_vus.vus_id = vus
        admin_vus.is_write = True
        db.session.add(admin_vus)
    for vus in vus_for_read:
        admin_vus = Admins_vuses()
        admin_vus.user_id = user.id
        admin_vus.vus_id = vus
        admin_vus.is_write = False
        db.session.add(admin_vus)
    db.session.commit()
    return gen_success(message={'status':'ok'})


######
@app.route('/post_add_vus', methods=['POST'])
def post_add_vus():
    data = request.form
    vus = VUS.query.filter_by(number=data['number'], code=data['code']).first()
    if vus:
        return gen_success(message={'status':'error', 'error' : u'Специальность была добавлена ранее'})
    vus = VUS()
    vus.number=data['number']
    vus.code=data['code']
    vus.name1=data['name1']
    vus.name2=data['name2']
    db.session.add(vus)
    db.session.commit()
    return gen_success(message={'status':'ok'})

@app.route('/create_accounts', methods=['POST'])
def create_accounts():
    if 'file' not in request.files:
        return gen_error('Файл не выбран')
    file = request.files['file']

    # if user does not select file, browser also
    # submit a empty part without filename
    if file.filename == '' or not file:
        return gen_error(u'Файл не выбран')
    if file.filename[-4:] != 'xlsx':
        return gen_error(u'Файл должен быть формата .xlsx')

    if 'vus' not in request.form:
        return gen_error('Выберите ВУС')
    if 'completionYear' not in request.form:
        return gen_error('Введите год поступления')

    completionYear = request.form['completionYear']
    if completionYear == '':
        return gen_error('Введите год поступления')

    vus = map(int, request.form['vus'].split())
    vus = VUS.query.filter_by(number=vus[0], code=vus[1]).first()
    if vus is None:
        return gen_error('Such vus not yet exists in this system')

    wb = load_workbook(file)
    active = wb.active
    userNames = User.query.with_entities(User.login)

    for idx, row in enumerate(active.rows, start = 1):
        login = ''

        #фамилия
        for char in row[0].value:
            login += vocabulary[char.lower()]
        login += u'.'

        #инициалы имени и отчества
        firstNameShort = row[1].value[0].lower()
        middleNameShort = row[2].value[0].lower()
        login += vocabulary[firstNameShort] + u'.'
        login += vocabulary[middleNameShort] + u'.'
        
        login += completionYear
        password = ''.join(random.choice(string.ascii_uppercase + string.ascii_lowercase + string.digits) for _ in range(8))

        for name in userNames:
            if login == name.login:
                return gen_error(u'В системе уже существует аккаунт: ' + login)

        info = {
            'lastName': row[0].value,
            'firstName': row[1].value,
            'middleName': row[2].value,
            'year': completionYear,
            'vus': vus
        }

        create_account(login, password, info)

        active.cell(row = idx, column = 4, value = login)
        active.cell(row = idx, column = 5, value = password)

    path = os.path.join(USER_PATH, 'logins.xlsx')
    wb.save(path)

    return gen_success(url = '/static/user_data/logins.xlsx')

#consent to procesing
@app.route('/consent_proc', methods=['POST'])
def consent_proc():
    current_user.processing_consent = True
    db.session.commit()
    return gen_success(message={'status':'ok'})

@app.route('/add_document', methods=['POST'])
def add_document():
    if 'file' not in request.files:
        return gen_error('No file sent')
    file = request.files['file']

    # if user does not select file, browser also
    # submit a empty part without filename
    if file.filename == '' or not file:
        return gen_error('No file selected')

    if 'name' not in request.form:
        return gen_error('No name document could not be created')
    if 'vus' not in request.form:
        return gen_error('No vus document could not be created')

    vus = map(int, request.form['vus'].split())
    vus = VUS.query.filter_by(number=vus[0], code=vus[1]).first()
    if vus is None:
        return gen_error('Such vus not yet exists in this system')

    docs = os.listdir(os.path.join(USER_PATH, 'documents'))
    filename = file.filename
    if filename in docs:
        i = 1
        while filename[:-5] + '_' + str(i) + filename[-5:] in docs:
            i += 1
        filename = filename[:-5] + '_' + str(i) + filename[-5:]

    file.save(os.path.join(USER_PATH, 'documents', filename))
    d = Document(name=request.form['name'], vus_id=vus.id, filename=filename)
    db.session.add(d)
    db.session.commit()

    return gen_success(filename=filename, message='Success!')

@app.route('/delete_document', methods=['POST'])
def delete_document():
    data = json.loads(request.data)
    docId = data['docId']

    document = Document.query.filter_by(id = docId).first()
    filePath = os.path.join(USER_PATH, 'documents', document.filename)


    if document is None:
        return gen_error('No document with such id')

    os.remove(filePath)

    Document.query.filter_by(id = docId).delete()
    db.session.commit()

    return gen_success(message='Success!')

### POSTs

def check_errors_in_input_data(table, data):
    tableclass = get_class_by_tablename(table)()
    errors = []
    for field, value in data.iteritems():
        if len(field) and not len(value):
            errors.append( u'Заполните поле "' + tableclass.get_russian_name( field ) + '"' )
    return errors

def save_not_fixed_section_information(data):
    user_id = data['user_id']

    elements = json.loads(data['elements'])
    # проверяем пустые поля
    errors = []
    for element in elements:
        errors += check_errors_in_input_data(data['table'], element)

    if len(errors):
        return gen_success(message = {'status':'error', 'errors':"<br>".join(errors)})

    student_info = User.query.get( user_id ).students_info
    records = student_info[data['table']]
    tableclass = get_class_by_tablename(data['table'])
    
    ### delete existing
    if records and len(records):
        for record in records:
            db.session.delete(record)

    ### add new records
    for element in elements:
        element_fields = { field : element[field] for field in element if hasattr(tableclass, field) }
        new_record = tableclass()
        for field, value in element_fields.iteritems():
            new_record[field] = value
        new_record.student_info    = student_info
        new_record.student_info_id = student_info.id
        db.session.add(new_record)
        db.session.commit()

    student_info['table_' + data['table']] = TABLE_STATES['EDITED']
    db.session.commit()

    return gen_success(message = {'status':'ok'} )
        

def save_section_information(data):
    
    # проверяем пустые поля
    if 'elements' in data:
        return save_not_fixed_section_information(data)

    user_id = data['user_id']
    errors = check_errors_in_input_data(data['table'], data)
    if len(errors):
        return gen_success(message = {'status':'error', 'errors' : "<br>".join(errors) })
    
    student_info = User.query.get( user_id ).students_info
    table = student_info[data['table']]
    # обновляем таблицу
    if table:
        tableclass = get_class_by_tablename(data['table'])
        table_fields = { field : data[field] for field in data if hasattr(tableclass, field) }
        for field, value in table_fields.iteritems():
            table[field] = value
        student_info['table_' + data['table']] = TABLE_STATES['EDITED']
        db.session.commit()

    return gen_success(message = {'status':'ok'} )

def approve_all_sections(data):
    user = User.query.get( int(data['user_id']) )
    student_info = user.students_info
    for table in get_user_tables():
        student_info['table_' + table] = TABLE_STATES['APPROVED']
    is_all_approved = True

    for table in (get_user_tables() + get_admin_tables()):
        if student_info['table_' + table] != TABLE_STATES['APPROVED']:
            is_all_approved = False
            break

    if is_all_approved:
        user.approved = True

    db.session.commit()
    return gen_success(message = {'status':'ok'} )

def change_vus_status(data):
    vus_id = int(data['vus_id'])
    new_status = int(data['new_status'])

    vus = VUS.query.get( vus_id )
    vus['is_active'] = new_status

    db.session.commit()
    return gen_success(message = {'status':'ok'} )

def change_section_state(data):
                    
    new_state = int(data['new_state'])
    user = User.query.get( int(data['user_id']) )
    student_info = user.students_info

    student_info['table_' + data['table']] = new_state
    if new_state == TABLE_STATES['DECLINED']:
        user.approved = False
        student_info['comments'][data['table'] + '_comment'] = data['comment'] if 'comment' in data else ''
    else:
        student_info['comments'][data['table'] + '_comment'] = ''

    is_all_approved = True
    for table in (get_user_tables() + get_admin_tables()):
        if student_info['table_' + table] != TABLE_STATES['APPROVED']:
            is_all_approved = False
            break

    if is_all_approved:
        user.approved = True

    db.session.commit()
    return gen_success(message = {'status':'ok'} )

def send_quiz_to_check(data):
    return gen_success(message =  {'status':'ok'})

def delete_user(user_id):
    user = User.query.get(user_id)
    student_info = user.students_info

    for table in (get_user_tables() + get_admin_tables()):
        section = get_class_by_tablename(table)()
        if (student_info[table]):
            if section.is_fixed:
                db.session.delete(student_info[table])
            else:
                for record in student_info[table]:
                    db.session.delete(record) 
    
    db.session.delete(student_info['comments'])
    db.session.delete(student_info)
    db.session.delete(user)
    db.session.commit()
    return True

def post_delete_user(data):
    user_id = int(data['user_id'])
    delete_user(user_id)
    return gen_success(message =  {'status':'ok'})

### SEARCH

def searchUsers(data): 
    sqlRequest = getSqlRequest(data['lastName'], data['year'], data['vus'], int(data['is_approved']))

    requestResult = db.engine.execute(sqlRequest)

    searchResult = []
    for row in requestResult:
        matchedUser = {
            'id' : row[0],
            'firstName' : row[6],
            'middleName' : row[7],
            'lastName' : row[3],
            'year' : row[2][-4:],
            'vus' : '%03d %03d' % (row[4], row[5]),
            'is_approved' : u'да' if row[8] else u'нет',
        }
        searchResult.append(matchedUser)
    
    searchResult = sorted(searchResult, key=( lambda u: (u['year'], u['vus'], u['lastName']) ))

    return gen_success(result = searchResult)

def getSqlRequest(lastName, year, vusStr, is_approved=False):
    
    sqlRequest = "select u_id, u_role, u_login, bi_last_name,\
        VUS.number as 'vus_num',\
        VUS.code as 'vus_code',\
        bi_first_name,\
        bi_middle_name,\
        u_is_approved \
        from (\
        select user.id as 'u_id',\
        user.role as 'u_role',\
        user.login as 'u_login',\
        user.vus_id as 'u_vus_id',\
        user.approved as 'u_is_approved',\
        bi_last_name,\
        bi_first_name,\
        bi_middle_name \
        from (\
        select\
        student_info.user_id as 'si_user_id',\
        basic_information.last_name as 'bi_last_name',\
        basic_information.first_name as 'bi_first_name',\
        basic_information.middle_name as 'bi_middle_name' \
        from student_info left join basic_information\
        on student_info.id = basic_information.student_info_id) as X\
        left join user\
        on X.si_user_id = user.id) as Y\
        left join VUS\
        on Y.u_vus_id = VUS.id "

    conds = ["u_role = '0'"]
    
    if is_approved > 0:
        conds.append('u_is_approved != 0')
    
    if lastName != '':
        conds.append("bi_last_name = '" + unicode(lastName) + "'")

    if year != '':
        conds.append("u_login like '%" + unicode(year) + "'")

    if vusStr != u'не выбрано':
        vus = map(int, vusStr.split())

        conds.append("vus_num = '" + str(vus[0]) + "'")
        conds.append("vus_code = '" + str(vus[1]) + "'")


    whereBlock = 'where ' + ' and '.join(conds) + ';'

    return text(sqlRequest + whereBlock)



### GENERATING DOCUMENTS

def parseDocument(doc, accessor):

    regex = re.compile('\{[a-zA-Z0-9_.@]+\}')
    
    parseParagraphs(doc.paragraphs, accessor, regex)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parseParagraphs(cell.paragraphs, accessor, regex)
                # для формы 12
                for table1 in cell.tables:
                    for row1 in table1.rows:
                        for cell1 in row1.cells:
                            parseParagraphs(cell1.paragraphs, accessor, regex)

def parseParagraphs(paragraphs, accessor, regex):
    
    for p in paragraphs:
        inline = p.runs
        keyString = ''
        started = False
        finished = False
        for i in range(len(inline)):
            #print p.text
            if len(inline[i].text) > 0 and inline[i].text[0] == '{':
                started = True
            if len(inline[i].text) > 0 and inline[i].text[-1] == '}':
                #print inline[i].text.rstrip()
                finished = True
            if started:
                keyString += inline[i].text
                inline[i].text = ''
            
            #if started:
                #print keyString, started, finished

            if started and finished:
                inline[i].text = keyString
                iteritems = regex.finditer(keyString)
                for item in iteritems:
                    key = item.group()
                    value = unicode(accessor[key])
                    value = value if value != None else u'НЕПРАВИЛЬНЫЙ КЛЮЧ!'
                    #print key, value, inline[i].text
                    text = inline[i].text.replace(key, value)
                    inline[i].text = text
                    #print inline[i].text
                started = False
                finished = False
                keyString = ''
                #print 'done'

def generateDocuments(data):
    userIDs = json.loads(data['userIDs'])
    docIDs = json.loads(data['docIDs'])
    #print data
    generationDate = data['generation_date']
    #print generationDate
    
    users = User.query.filter(User.id.in_(userIDs)).all()
    documents = Document.query.filter(Document.id.in_(docIDs)).all()
    vuses = VUS.query.all()

    basedir = os.path.join(USER_PATH, 'documents', 'temp')
    if not os.path.exists(basedir):
        os.makedirs(basedir)

    if not users:
        return gen_success(success = False, message = 'Выберите хотя бы одного пользователя')
    if not documents:
        return gen_success(success = False, message = 'Выберите хотя бы один документ')
    
    for user in users:
        vus = [vus for vus in vuses if vus.id == user.vus_id][0]
        accessor = Students_info_lables_accessor(user.students_info, vus, generationDate)
        for document in documents:
            docPath = os.path.join(USER_PATH, 'documents', document.filename)
            doc = Doc(docPath)
            
            parseDocument(doc, accessor)
                
            separators = [pos for pos, char in enumerate(document.filename) if char == '_']
            doc_name = document.filename[:-5] if len(separators) == 0 else document.filename[:separators[-1]]    
            doc_name = os.path.join(USER_PATH, 'documents', 'temp', doc_name + '_' + user.login + '.docx')
            doc.save(doc_name)
    
    zippath = os.path.join(USER_PATH, 'Documents.zip')
    zipf = ZipFile(zippath, 'w', ZIP_DEFLATED)
    
    for root, dirs, files in os.walk(basedir):
        for fn in files:
            absfn = os.path.join(root, fn)
            zfn = absfn[len(basedir)+len(os.sep):]
            zipf.write(absfn, zfn)
            os.remove(absfn)
    zipf.close()

    return gen_success(url = '/static/user_data/Documents.zip', success = True)



@app.route('/post_query', methods=['POST'])
def post_query():
    data = request.form
    try:
        if 'do' in data and data['do'] in POST_METHODS:
            return POST_METHODS[data['do']](data)
        else:
            return gen_success(message = {'status':'error', 'error':'post method not defined'})
    except Exception as err:
        return gen_success(message = {'status':'error', 'error':'error in post method ' + str(err)})


# list of post methods
POST_METHODS = dict( [ (table, save_section_information) for table in (get_user_tables() + get_admin_tables()) ] )
POST_METHODS.update( {
                'searchUsers': searchUsers,
                'generateDocuments': generateDocuments,
                'send_quiz_to_check': send_quiz_to_check,
                'change_section_state' : change_section_state,
                'approve_all_sections' : approve_all_sections,
                'delete_user' : post_delete_user,
                'change_vus_status' : change_vus_status,

                })


